from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.utils.encoding import escape_uri_path
from wsgiref.util import FileWrapper
from pool.models import Category, Subject, Chapter
from pool.views import class_list, sub_class_list, type_sc_abbr, type_sc_full, type_en_abbr
from .models import *
from control.models import User
from PIL import Image
from zipfile import ZipFile
from io import BytesIO
from dateutil import parser
from Crypto.Cipher import AES
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
import json
import re
import uuid
import xlwt
import random
import string
import base64


# Create your views here.
@login_required()
def exams(request):
    operator = User.objects.get(username=request.user.username)
    access = operator.profile.access
    if not operator.profile.view_access.get('exam'):
        return redirect('home:dashboard')
    chapter_points = tuple(map(lambda num: 'point{}'.format(num), range(1, 8)))
    chapter_difficulties = tuple(map(lambda num: 'difficulty{}'.format(num), range(1, 8)))
    chapter_q_nums = tuple(map(lambda num: 'q_num{}'.format(num), range(1, 8)))
    context = {
        'all_categories': Category.objects.all().order_by('index'),
        'all_subjects': Subject.objects.all().order_by('index'),
        'all_chapters': Chapter.objects.all().order_by('index'),
        'all_strategies': Strategy.objects.all().order_by('index'),
        'all_ads': Advertisement.objects.all(),
        'all_agreements': Agreement.objects.all(),
        'all_exams': Exam.objects.all().order_by('-created'),
        'all_student_lists': StudentList.objects.all(),
        'type_sc_abbr': type_sc_abbr,
        'chapter_info': tuple(zip(type_sc_abbr, chapter_points, chapter_difficulties, chapter_q_nums)),
        'can_manage_strategy': access.get('manage_strategy', False),
        'can_generate_exam': access.get('generate_exam', False),
        'can_download_exam': access.get('download_exam', False),
        'can_manage_exam': access.get('manage_exam', False),
    }
    if request.method == 'POST':
        if 'add_strategy' in request.POST:
            new_object = Strategy(
                subject_id=int(request.POST.get('strategy_subject')),
                name=request.POST.get('strategy_name'),
                index=request.POST.get('strategy_index'),
            )
            new_object.description = request.POST.get('strategy_description')
            if request.POST.get('strategy_timer'):
                new_object.timer = int(request.POST.get('strategy_timer'))
            new_object.save()
        elif 'sample_check' in request.POST:
            this_object = Strategy.objects.get(id=int(request.POST.get('selected_strategy')))
            return generate_sample(this_object.plan, str(this_object))
        elif 'strategy_plan' in request.POST:
            this_object = Strategy.objects.get(id=int(request.POST.get('selected_strategy')))
            plan = json.loads(request.POST.get('strategy_plan'))
            this_object.plan = plan
            this_object.save()
        elif 'edit_strategy' in request.POST or 'duplicate_strategy' in request.POST:
            this_object = Strategy.objects.get(id=int(request.POST.get('strategy_id')))
            this_object.subject_id = int(request.POST.get('strategy_subject'))
            this_object.name = request.POST.get('strategy_name')
            this_object.index = request.POST.get('strategy_index')
            this_object.description = request.POST.get('strategy_description')
            if request.POST.get('strategy_timer'):
                this_object.timer = int(request.POST.get('strategy_timer'))
            else:
                this_object.timer = 100
            if 'duplicate_strategy' in request.POST:
                this_object.pk = None
            this_object.save()
        elif 'delete_strategy' in request.POST:
            this_object = Strategy.objects.get(id=int(request.POST.get('strategy_id')))
            this_object.delete()
        elif 'add_ad' in request.POST:
            new_object = Advertisement(name=request.POST.get('ad_name'), image=request.FILES.get('ad_image'))
            new_object.description = request.POST.get('ad_description')
            new_object.save()
        elif 'add_agreement' in request.POST:
            new_object = Agreement(name=request.POST.get('agreement_name'),
                                   image=request.FILES.get('agreement_image'))
            new_object.description = request.POST.get('agreement_description')
            new_object.save()
        elif 'edit_ad' in request.POST:
            this_object = Advertisement.objects.get(id=int(request.POST.get('ad_id')))
            if request.POST.get('ad_name'):
                this_object.name = request.POST.get('ad_name')
            if request.POST.get('ad_changed'):
                this_object.image = request.FILES.get('ad_image')
            this_object.description = request.POST.get('ad_description')
            this_object.save()
        elif 'edit_agreement' in request.POST:
            this_object = Agreement.objects.get(id=int(request.POST.get('agreement_id')))
            if request.POST.get('agreement_name'):
                this_object.name = request.POST.get('agreement_name')
            if request.POST.get('agreement_changed'):
                this_object.image = request.FILES.get('agreement_image')
            this_object.description = request.POST.get('agreement_description')
            this_object.save()
        elif 'delete_ad' in request.POST:
            this_object = Advertisement.objects.get(id=int(request.POST.get('ad_id')))
            this_object.delete()
        elif 'delete_agreement' in request.POST:
            this_object = Agreement.objects.get(id=int(request.POST.get('agreement_id')))
            this_object.delete()
        elif ('add_exam' in request.POST) or ('add_hw' in request.POST):
            title = request.POST.get('exam_title')
            location = request.POST.get('exam_location')
            section = request.POST.get('exam_section')
            date = request.POST.get('exam_date')
            agreement = request.POST.get('exam_agreement')
            ad = request.POST.get('exam_ad')
            # noinspection PyBroadException
            try:
                dt = parser.parse(date)
                date = '{}{:0>2}{:0>2}'.format(dt.year, dt.month, dt.day)
            except Exception:
                messages.error(request, '日期格式错误！')
                return redirect('exam:exams')
            this_strategy = Strategy.objects.get(id=int(request.POST.get('exam_strategy_id')))
            plan = this_strategy.plan
            # Parse students and photos
            students = request.POST.get('students').split('\r\n')
            student_photos = request.FILES.getlist('exam_photos')
            fvr = validate_students(students, student_photos)
            if not fvr[0]:
                messages.error(request, fvr[1])
                return redirect('exam:exams')
            student_json, student_photos_dict = fvr[2]
            # Parse strategy plans
            exam_set = get_exam_set(plan, student_json)
            if not exam_set[0]:
                messages.success(request, exam_set[1])
                return redirect('exam:exams')
            # Prepare the zip file
            buffer_excel = generate_excel({'student_json': student_json, }, True)
            buffer_zf = generate_zip(student_photos_dict, 'exams', {
                'problems': exam_set[2],
                'students': exam_set[3],
                'agreement': int(agreement),
                'ad': int(ad),
                'basics': {
                    'title': title,
                    'location': location,
                    'section': section,
                    'date': date,
                    'timer': this_strategy.timer,
                    'subject': str(this_strategy.subject),
                    'is_homework': 'add_hw' in request.POST
                }
            })
            # Save objects to database
            new_object = Exam(title=title, location=location, section=section, date=date, plan=plan)
            new_object.learning = 'add_hw' in request.POST
            new_object.package.save('{}.zip'.format(uuid.uuid4().hex), ContentFile(buffer_zf))
            new_student_list_file = StudentListFile(exam_id=new_object.id)
            new_student_list_file.student_list.save('{}.xls'.format(uuid.uuid4().hex), ContentFile(buffer_excel))
            if request.POST.get('student_list_name') != '':
                new_student_list = StudentList(
                    name=request.POST.get('student_list_name'),
                    student_list=student_json
                )
                new_student_list.save()
        messages.success(request, '操作成功！')
        return redirect('exam:exams')
    return render(request, 'exam/exams.html', context)


@login_required()
def certification(request):
    operator = User.objects.get(username=request.user.username)
    access = operator.profile.access
    if not operator.profile.view_access.get('certification'):
        return redirect('home:dashboard')
    context = {
        'all_certifications': GovernmentCertification.objects.all().order_by('-created'),
        'all_exams': Exam.objects.all(),
        'all_student_lists': StudentList.objects.all(),
        'all_code_categories': CodeCategory.objects.all().order_by('index'),
        'all_strategies': Strategy.objects.all().order_by('index'),
        'all_code_subjects': CodeSubject.objects.all().order_by('code'),
        'can_download_certification': access.get('download_certification', False),
        'can_manage_certification': access.get('manage_certification', False),
        'can_manage_code': access.get('manage_code', False),
    }
    if request.method == 'POST':
        if 'btn_certification' in request.POST:
            new_object = GovernmentCertification(
                name=request.POST.get('certification_name'),
                project=request.POST.get('certification_project'),
                verified=request.POST.get('certification_verified'),
                school=request.POST.get('certification_school')
            )
            new_object_subject = request.POST.get('certification_subject')
            if new_object_subject is None:
                new_object_subject = ''
            new_object.subject = new_object_subject
            # Parse students and photos
            students = request.POST.get('students').split('\r\n')
            student_photos = request.FILES.getlist('certification_photos')
            fvr = validate_students(students, student_photos, repeated_id=True)
            if not fvr[0]:
                messages.error(request, fvr[1])
                return redirect('exam:certification')
            student_json, student_photos_dict = fvr[2]
            # Prepare the zip file
            buffer_student_list = generate_excel({
                'student_json': student_json,
                'name': new_object.name,
                'project': new_object.project,
                'verified': new_object.verified,
                'school': new_object.school,
                'subject': new_object.subject,
            })
            buffer_zf = generate_zip(student_photos_dict, 'certification', {
                'excel': ('{}.xls'.format(new_object.name if new_object.name else '未命名表单'), buffer_student_list)
            })
            new_object.student_list = student_json
            new_object.package.save('{}.zip'.format(uuid.uuid4().hex), ContentFile(buffer_zf))
            if request.POST.get('student_list_name') != '':
                new_student_list = StudentList(
                    name=request.POST.get('student_list_name'),
                    student_list=student_json
                )
                new_student_list.save()
        elif 'add_code' in request.POST:
            code_category_id = int(request.POST.get('code_category'))
            new_object = CodeSubject(
                name=request.POST.get('code_name'),
                code=request.POST.get('code_code')
            )
            if code_category_id == 0:  # Create new
                new_code_category = CodeCategory(
                    name=request.POST.get('new_code_category'),
                    index=request.POST.get('new_code_category_index')
                )
                new_code_category.save()
                new_object.category = new_code_category
            else:
                new_object.category_id = code_category_id
            if request.POST.get('code_price'):
                new_object.price = float(request.POST.get('code_price'))
            if request.POST.get('code_description'):
                new_object.description = request.POST.get('code_description')
            new_object.save()
        elif 'add_category' in request.POST:
            new_code_category = CodeCategory(
                name=request.POST.get('category_name'),
                index=request.POST.get('category_index')
            )
            new_code_category.save()
        elif 'edit_code' in request.POST:
            code_category_id = int(request.POST.get('code_category'))
            this_object = CodeSubject.objects.get(id=int(request.POST.get('code_id')))
            if code_category_id == 0:  # Create new
                new_code_category = CodeCategory(
                    name=request.POST.get('new_code_category'),
                    index=request.POST.get('new_code_category_index')
                )
                new_code_category.save()
                this_object.category = new_code_category
            else:
                this_object.category_id = code_category_id
            if request.POST.get('code_price'):
                this_object.price = float(request.POST.get('code_price'))
            else:
                this_object.price = 0
            this_object.description = request.POST.get('code_description')
            this_object.save()
        elif 'edit_category' in request.POST:
            this_object = CodeCategory.objects.get(id=int(request.POST.get('category_id')))
            this_object.name = request.POST.get('category_name')
            this_object.index = request.POST.get('category_index')
            this_object.save()
        elif 'delete_code' in request.POST:
            this_object = CodeSubject.objects.get(id=int(request.POST.get('code_id')))
            this_object.delete()
        elif 'delete_category' in request.POST:
            this_object = CodeCategory.objects.get(id=int(request.POST.get('category_id')))
            this_object.delete()
        messages.success(request, '操作成功！')
        return redirect('exam:certification')
    return render(request, 'exam/certification.html', context)


def get_exam(request, exam_id):
    this_file = Exam.objects.get(id=exam_id)
    file_path = 'media/{}'.format(this_file.package.name)
    file_name = '-'.join([this_file.title, this_file.location, this_file.section, this_file.date, ])
    response = HttpResponse(FileWrapper(open(file_path, 'rb')), content_type='application/zip')
    response['Content-Disposition'] = "attachment; filename*=utf-8''{}.exam".format(escape_uri_path(file_name))
    return response


@login_required()
def delete_exam(request, exam_id):
    this_object = Exam.objects.get(id=exam_id)
    this_object.delete()
    return redirect('exam:exams')


def get_certification(request, certification_id):
    this_file = GovernmentCertification.objects.get(id=certification_id)
    file_path = 'media/{}'.format(this_file.package.name)
    file_name = this_file.name if this_file.name else '未命名认证包'
    response = HttpResponse(FileWrapper(open(file_path, 'rb')), content_type='application/zip')
    response['Content-Disposition'] = "attachment; filename*=utf-8''{}.zip".format(escape_uri_path(file_name))
    return response


@login_required()
def delete_certification(request, certification_id):
    this_object = GovernmentCertification.objects.get(pk=certification_id)
    this_object.delete()
    return redirect('exam:certification')


def get_exam_list(request, exam_id):
    this_exam = Exam.objects.get(id=exam_id)
    this_file = StudentListFile.objects.get(exam_id=exam_id)
    file_path = 'media/{}'.format(this_file.student_list.name)
    file_name = '-'.join([this_exam.title, this_exam.location, this_exam.section, this_exam.date, ])
    response = HttpResponse(FileWrapper(open(file_path, 'rb')), content_type='application/vnd.ms-excel')
    response['Content-Disposition'] = "attachment; filename*=utf-8''{}.xls".format(escape_uri_path(file_name))
    return response


@login_required()
def delete_student_list(request, list_id):
    this_object = StudentList.objects.get(id=list_id)
    this_object.delete()
    return redirect('exam:exams')


def wrong_message(warning):
    return False, warning, {}


def validate_students(students, photos, repeated_id=False):
    student_json = {}
    student_photos_dict = {}
    for student in students:
        student = student.strip()
        if student:
            r = re.match('(.*)-(.*)-(.*)', student)
            if r:
                this_exam_id = r.group(1).strip()
                this_student_name = r.group(2).strip()
                this_student_id = r.group(3).strip()
                if (not repeated_id) and (this_exam_id in student_json):
                    return wrong_message('您有重复的考试号！考试号：{}'.format(this_exam_id))
                generalized_exam_id = student if repeated_id else this_exam_id
                student_json[generalized_exam_id] = (this_student_name, this_student_id, this_exam_id)
                student_photos_dict['{}{}'.format(this_exam_id, this_student_name)] = None
            else:
                return wrong_message('学生列表输入格式有误！位置：{}'.format(student))
    for photo in photos:
        photo_name = photo.name.strip()
        r = re.match('(.*)\.(.*)', photo_name)
        if r:
            id_name = digest_id_name(r.group(1).strip())
            if id_name not in student_photos_dict:
                return wrong_message('您有多余的学生照片，文件名为{}'.format(photo_name))
            if student_photos_dict[id_name] is not None:
                return wrong_message('您有重复学生照片，文件名为{}'.format(photo_name))
            student_photos_dict[id_name] = photo
        else:
            return wrong_message('无法识别的文件，文件名为{}'.format(photo_name))
    return True, '', (student_json, student_photos_dict,)


def digest_id_name(id_name):
    r = re.match('(.*)-(.*)', id_name)
    if r:
        id_name = '{}{}'.format(r.group(1).strip(), r.group(2).strip())
    r = re.match('(.*)[\t\s]+(.*)', id_name)
    if r:
        id_name = '{}{}'.format(r.group(1).strip(), r.group(2).strip())
    r = re.match('(\d+)(.*)', id_name)
    if r:
        id_name = '{}{}'.format(int(r.group(1).strip()), r.group(2).strip())
    return id_name


def get_exam_set(plan, students):
    # Get all questions, skip 0 plan selection
    all_problems = {}
    for plan_line in plan:
        chapter = Chapter.objects.get(pk=plan_line[0])
        chapter_set = {}
        for i in range(7):
            if plan_line[i + 1][0] > 0:
                chapter_problem_set = getattr(chapter, '{}_set'.format(class_list[i].__name__.lower())).all()
                chapter_set.update({type_en_abbr[i]: chapter_problem_set, })
            else:
                chapter_set.update({type_en_abbr[i]: None, })
        all_problems.update({plan_line[0]: chapter_set, })
    # Form random exams for each student and append questions to selected set
    problems = {}
    student_info = {}
    for exam_id in students:
        my_problems = []
        for plan_line in plan:
            my_chapter = {plan_line[0]: {}}
            for i in range(7):
                potential_problems = all_problems[plan_line[0]][type_en_abbr[i]]
                num_of_problems, problem_points = plan_line[i + 1]
                if potential_problems:
                    selected_problems = []
                    if len(potential_problems) < num_of_problems:
                        return wrong_message('选择的题量超出录入量，请检查！')
                    for j in random.sample(range(len(potential_problems)), num_of_problems):
                        selected_problems.append(potential_problems[j])
                        # Get problem detail and update general problem set (aka problems)
                        if not type_en_abbr[i] in problems:
                            problems[type_en_abbr[i]] = {}
                        if not potential_problems[j].id in problems[type_en_abbr[i]]:
                            problem_detail = get_problem_details(type_en_abbr[i], potential_problems[j].id,
                                                                 problem_points)
                            problems[type_en_abbr[i]][potential_problems[j].id] = problem_detail
                    # Sort selected problems based on their index
                    selected_problems.sort(key=lambda x: x.index, reverse=False)
                    my_chapter[plan_line[0]][type_en_abbr[i]] = list(map(lambda x: x.id, selected_problems))
            my_problems.append(my_chapter)
        student_info[exam_id] = {
            'name': students[exam_id][0],
            'student_id': students[exam_id][1],
            'problems': my_problems,
            'key': ''.join(random.choices(string.ascii_letters + string.digits, k=16))
        }
    return True, '', problems, student_info


def generate_sample(plan, strategy_name):
    exam_set = get_exam_set(plan, {'SAMPLE': ['S_NAME', 'S_ID']})
    problems = exam_set[2]
    problem_list = exam_set[3]['SAMPLE']['problems']
    buffer_zf = BytesIO()
    zf = ZipFile(buffer_zf, 'w')
    zf.writestr('样卷.docx', generate_word(problems, problem_list))
    zf.writestr('样卷答案.docx', generate_word(problems, problem_list, True))
    zf.close()
    buffer_zf.seek(0)
    response = HttpResponse(buffer_zf.read(), content_type='application/zip')
    response['Content-Disposition'] = "attachment; filename*=utf-8''{}.zip".format(escape_uri_path(strategy_name))
    buffer_zf.close()
    return response


def generate_word(problems, problem_list, is_answer_sheet=False):
    buffer_word = BytesIO()
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    word_content = []
    # Unpack problem list
    index = 0
    for chapter in problem_list:
        for chapter_id in chapter:
            for problem_type in chapter[chapter_id]:
                for problem_id in chapter[chapter_id][problem_type]:
                    index += 1
                    content = problems[problem_type][problem_id]
                    # Shared parts - Header and Body
                    if problem_type == 'cp' and index > 1:
                        word_content.append(None)
                    word_content.append([f'第{index}题 {content["type_sc"]}（{content["points"]}分）', 'bold'])
                    word_content.extend(content['desc_lines'])
                    if 'image' in content:
                        word_content.append([content['image'], 'image'])
                    if problem_type == 'cp':
                        word_content.append('')
                        for sub in content['sub']:
                            sub_points = round(content['points'] * sub['percentage'] / 100)
                            word_content.append([f'第{sub["order"]}小题 {sub["type_sc"]}（{sub_points}分）', 'bold'])
                            word_content.extend(sub['desc_lines'])
                            if 'image' in sub:
                                word_content.append([sub['image'], 'image'])
                            # Body cont'd
                            if 'choice_lines' in sub:
                                for i in range(len(sub['choice_lines'])):
                                    word_content.append(sub['choice_lines'][i])
                                    if chr(65 + i) in sub:
                                        word_content.append([sub[chr(65 + i)], 'image'])
                            # Footer
                            ans_lines = '\n'.join(sub['ans_lines'])
                            word_content.append(f'答案：{ans_lines if is_answer_sheet else "___________"}')
                            if is_answer_sheet and 'answer_image' in sub:
                                word_content.append([sub['answer_image'], 'image'])
                            word_content.append('')
                        word_content.append(None)
                    else:
                        # Body cont'd
                        if 'choice_lines' in content:
                            for i in range(len(content['choice_lines'])):
                                word_content.append(content['choice_lines'][i])
                                if chr(65 + i) in content:
                                    word_content.append([content[chr(65 + i)], 'image'])
                        # Footer
                        ans_lines = '\n'.join(content['ans_lines'])
                        if 'error' in content:
                            ans_lines += f'（±{content["error"]}%）'
                        word_content.append(f'答案：{ans_lines if is_answer_sheet else "___________"}')
                        if is_answer_sheet and 'answer_image' in content:
                            word_content.append([content['answer_image'], 'image'])
                        word_content.append('')
    word_content = word_content if word_content[-1] is not None else word_content[:-1]
    on_content = False
    for paragraph in word_content:
        if paragraph is None:
            if on_content:
                on_content = False
                document.add_page_break()
        elif isinstance(paragraph, list):
            on_content = True
            if paragraph[1] == 'bold':
                p = document.add_paragraph()
                p.add_run(paragraph[0]).bold = True
                p.style = document.styles['Normal']
            elif paragraph[1] == 'image':
                p = document.add_paragraph()
                run = p.add_run()
                inline_shape = run.add_picture(paragraph[0])
                if inline_shape.width > Inches(5):
                    inline_shape.height = Inches(5 * inline_shape.height / inline_shape.width)
                    inline_shape.width = Inches(5)
        else:
            on_content = True
            p = document.add_paragraph(paragraph)
            p.style = document.styles['Normal']
    document.save(buffer_word)
    buffer_value = buffer_word.getvalue()
    buffer_word.close()
    return buffer_value


def generate_excel(form_data, is_sign_sheet=False):
    buffer_excel = BytesIO()
    student_json = form_data['student_json']
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('考生名单')
    row_num = 0
    font_style = xlwt.XFStyle()
    # First the header
    headers = [
        '序号', '姓名', '考试项目', '课程/级别', '通过方式',
        '证书编号', '身份证号', '成绩', '考试时间', '学校',
    ]
    if is_sign_sheet:
        headers = ['考试编号', '考生姓名', '证件号', '签名']
    for col_num in range(len(headers)):
        ws.write(row_num, col_num, headers[col_num], font_style)
        ws.col(col_num).width = 256 * 20
    # Now comes the info
    for json_key in student_json:
        json_value = list(student_json[json_key])
        student_name = json_value[0]
        student_id = json_value[1]
        exam_id = json_key if len(json_value) == 2 else json_value[2]
        student_info = [exam_id, student_name, student_id, ]
        if not is_sign_sheet:
            student_info = [
                exam_id, student_name, form_data['project'], form_data['subject'], form_data['verified'],
                '', student_id, '', '', form_data['school'],
            ]
        row_num += 1
        for col_num in range(len(student_info)):
            ws.write(row_num, col_num, student_info[col_num], font_style)
    wb.save(buffer_excel)
    buffer_value = buffer_excel.getvalue()
    buffer_excel.close()
    return buffer_value


def generate_zip(photos, caller, other_files):
    buffer_zf = BytesIO()
    zf = ZipFile(buffer_zf, 'w')
    for id_name in photos:
        photo = photos[id_name]
        if photo:
            buffer_photo = BytesIO()
            try:
                pil_image = Image.open(photo)
                # x, y = pil_image.size
                # pil_image = pil_image.resize((160, 160 * y // x), Image.ANTIALIAS)
                pil_image = pil_image.convert('RGB')
                # pil_image.save(buffer_photo, format='JPEG', optimize=True, quality=95)
                pil_image.save(buffer_photo, format='JPEG')
                zf.writestr('photos/{}.jpg'.format(id_name), buffer_photo.getvalue())
            finally:
                buffer_photo.close()
    if caller == 'certification':
        my_excel = other_files.get('excel')
        zf.writestr(my_excel[0], my_excel[1])
    elif caller == 'exams':
        # key and iv
        key = iv = 'WXBR.1683.dev-30'
        # student.json
        buffer_students_json = BytesIO(get_encrypted_buffer(json.dumps(other_files.get('students')), key, iv))
        zf.writestr('students.json', buffer_students_json.getvalue())
        # Files
        problems = other_files.get('problems')
        for problem_type in problems:
            for problem_id in problems[problem_type]:
                this_problem = problems[problem_type][problem_id]
                media_attributes = ['image', 'attachment', 'answer_image', ]
                if 'choice_image_keys' in this_problem:
                    media_attributes.extend(this_problem['choice_image_keys'])
                for media in media_attributes:
                    if media in this_problem:
                        ext = '.'.join(this_problem[media].split('.')[1:])
                        archive_name = '{}-{}-{}.{}'.format(problem_type, problem_id, media, ext)
                        with open(this_problem[media], 'rb') as f:
                            buffer_file = BytesIO(get_encrypted_buffer(f.read(), key, iv))
                            zf.writestr('files/{}'.format(archive_name), buffer_file.getvalue())
                        this_problem[media] = archive_name
                if problem_type == 'cp':
                    for sub in problems[problem_type][problem_id]['sub']:
                        media_attributes = ['image', 'attachment', 'answer_image', ]
                        if 'choice_image_keys' in sub:
                            media_attributes.extend(sub['choice_image_keys'])
                        for media in media_attributes:
                            if media in sub:
                                ext = '.'.join(sub[media].split('.')[1:])
                                archive_name = 'sub{}-{}-{}.{}'.format(sub['type_en'], sub['id'], media, ext)
                                with open(sub[media], 'rb') as f:
                                    buffer_file = BytesIO(get_encrypted_buffer(f.read(), key, iv))
                                    zf.writestr('files/{}'.format(archive_name), buffer_file.getvalue())
                                sub[media] = archive_name
        # problems.json
        buffer_problems_json = BytesIO(get_encrypted_buffer(json.dumps(other_files.get('problems')), key, iv))
        zf.writestr('problems.json', buffer_problems_json.getvalue())
        # basics.json
        buffer_basics_json = BytesIO(get_encrypted_buffer(json.dumps(other_files.get('basics')), key, iv))
        zf.writestr('basics.json', buffer_basics_json.getvalue())
        # ad and agreement
        ad_id = other_files.get('ad')
        agreement_id = other_files.get('agreement')
        ad_path = 'static/img/ad_sc.jpg'
        agreement_path = 'static/img/agreement_sc.png'
        if ad_id != 0:
            ad_object = Advertisement.objects.get(pk=ad_id)
            ad_path = 'media/{}'.format(str(ad_object.image))
        if agreement_id != 0:
            agreement_object = Agreement.objects.get(pk=agreement_id)
            agreement_path = 'media/{}'.format(str(agreement_object.image))
        ad_ext = ad_path.split('.')[-1]
        agreement_ext = agreement_path.split('.')[-1]
        zf.write(ad_path, 'ad.{}'.format(ad_ext))
        zf.write(agreement_path, 'agreement.{}'.format(agreement_ext))
    zf.close()
    return buffer_zf.getvalue()


def get_problem_details(problem_type, problem_id, problem_points, is_sub=False):
    type_index = type_en_abbr.index(problem_type)
    problem_class = sub_class_list[type_index] if is_sub else class_list[type_index]
    this_problem = problem_class.objects.get(id=problem_id)
    result = {
        'id': problem_id,
        'type_en': problem_type,
        'type_sc': type_sc_full[type_index],
        'desc_lines': str(this_problem).split('\r\n')
    }
    if not is_sub:
        result['points'] = problem_points
    if type_index != 6:
        if type(this_problem.answer) == list:
            result['ans_lines'] = ''.join(this_problem.answer).split('\r\n')
        elif not this_problem.answer:
            result['ans_lines'] = '略。'
        else:
            result['ans_lines'] = '{}'.format(this_problem.answer).split('\r\n')
        if type_index == 0 or type_index == 1:
            choice_lines = []
            for i in range(len(this_problem.choices)):
                choice = this_problem.choices[i]
                choice_lines.append('(' + chr(65 + i) + ')' + choice)
            result['choice_lines'] = choice_lines
    else:  # Comprehensive part
        result_sub = []
        all_subs = []
        for sub_class in sub_class_list:
            sub_class_set = getattr(this_problem, '{}_set'.format(sub_class.__name__.lower())).all()
            all_subs.extend(list(sub_class_set))
        all_subs.sort(key=lambda x: x.order)
        for this_sub in all_subs:
            sub_index = sub_class_list.index(this_sub.__class__)
            result_sub.append(get_problem_details(type_en_abbr[sub_index], this_sub.id, 0, is_sub=True))
        result['sub'] = result_sub
    info_attributes = ('student_upload', 'need_answer', 'error', 'percentage', 'order')
    for info in info_attributes:
        if hasattr(this_problem, info):
            result[info] = getattr(this_problem, info)
    result.update(get_media_set(this_problem))
    return result


def get_media_set(this_problem):
    media_set = {}
    media_set_names = ('image', 'attachment', 'answer_image',)
    for media in media_set_names:
        if hasattr(this_problem, media):
            if getattr(this_problem, media):
                media_set[media] = 'media/{}'.format(str(getattr(this_problem, media)))
    image_set = this_problem.__class__.__name__.lower() + 'image_set'
    if hasattr(this_problem, image_set):
        choice_image_keys = []
        for choice_image in getattr(this_problem, image_set).all():
            media_set[choice_image.choice] = 'media/{}'.format(str(choice_image.image))
            choice_image_keys.append(choice_image.choice)
        media_set['choice_image_keys'] = choice_image_keys
    return media_set


def get_encrypted_buffer(raw_data, key, iv):
    def _pad(s):
        padding = AES.block_size - len(s) % AES.block_size
        return s + (padding * chr(padding)).encode()

    text_bytes = raw_data
    if not (type(raw_data) is bytes):
        text_bytes = raw_data.encode('utf-8')

    cipher = AES.new(key, AES.MODE_CBC, iv)
    msg = base64.b64encode(cipher.encrypt(_pad(text_bytes)))

    return msg


# AJAX functions
def change_category(request, category_id):
    if category_id == 0:
        subjects = Subject.objects.all()
    else:
        selected_category = Category.objects.get(pk=category_id)
        subjects = Subject.objects.filter(category=selected_category)
    subjects = subjects.order_by('index')
    subject_list = []
    strategy_list = []
    for sub in subjects:
        subject_list.append((sub.id, str(sub)))
        strategies = Strategy.objects.filter(subject=sub).order_by('index')
        for st in strategies:
            strategy_list.append((st.id, str(st)))
    return JsonResponse({'subjects': subject_list, 'strategies': strategy_list})


def change_subject(request, category_id, subject_id):
    if subject_id == 0:
        return change_category(request, category_id)
    else:
        strategies = Strategy.objects.filter(subject_id=subject_id)
    strategies = strategies.order_by('index')
    strategy_list = []
    for st in strategies:
        strategy_list.append((st.id, str(st)))
    return JsonResponse({'strategies': strategy_list})


def get_strategy(request, strategy_id):
    this_object = Strategy.objects.get(id=strategy_id)
    this_parent = this_object.subject
    this_grandparent = this_object.subject.category
    plan = []
    plan_matrix = this_object.plan
    if plan_matrix is not None:
        for plan_list in plan_matrix:
            plan.append({
                'plan_list': plan_list,
                'chapter_name': str(Chapter.objects.get(id=plan_list[0])),
                'list_total_points': sum([x * y for x, y in plan_list[1:]]),
            })
    res = {
        'category': this_grandparent.id,
        'subject': this_parent.id,
        'strategy': this_object.id,
        'name': this_object.name,
        'index': this_object.index,
        'full_path': '/'.join((this_grandparent.name, this_parent.name, this_object.name)),
        'full_index': '/'.join((this_grandparent.index, this_parent.index, this_object.index)),
        'description': this_object.description,
        'timer': '{}分钟'.format(this_object.timer),
        'timer_num': this_object.timer,
        'plan': plan,
    }
    return JsonResponse(res)


def get_picture(request, picture_type, picture_id):
    aa_dict = {
        'ad': ('/static/img/ad_sc.jpg', Advertisement,),
        'agreement': ('/static/img/agreement_sc.png', Agreement,),
    }
    if picture_id == 0:
        return JsonResponse({'image_path': aa_dict[picture_type][0]})
    this_object = aa_dict[picture_type][1].objects.get(id=picture_id)
    return JsonResponse({'image_path': '/media/' + str(this_object.image)})


def change_code_category(request, category_id):
    code_subjects = CodeSubject.objects.all() if category_id == 0 else CodeSubject.objects.filter(
        category_id=category_id)
    return JsonResponse({
        'code_subjects': list(map(lambda x: (x.id, str(x)), code_subjects))
    })


def get_code_category(request, category_id):
    this_object = CodeCategory.objects.get(id=category_id)
    return JsonResponse({
        'name': this_object.name,
        'index': this_object.index,
    })


def get_code(request, code_id):
    this_object = CodeSubject.objects.get(id=code_id)
    return JsonResponse({
        'category': this_object.category_id,
        'category_info': str(this_object.category),
        'name': this_object.name,
        'code': this_object.code,
        'subject_info': str(this_object),
        'price': this_object.price,
        'description': this_object.description
    })


def get_student_list(request, list_id):
    this_object = StudentList.objects.get(id=list_id)
    return JsonResponse({
        'student_list': this_object.student_list,
    })


def get_ad(request, ad_id):
    this_object = Advertisement.objects.get(id=ad_id)
    return JsonResponse({
        'name': this_object.name,
        'description': this_object.description,
        'image': str(this_object.image).split('/')[-1],
        'image_path': '/media/' + str(this_object.image),
    })


def get_agreement(request, agreement_id):
    this_object = Agreement.objects.get(id=agreement_id)
    return JsonResponse({
        'name': this_object.name,
        'description': this_object.description,
        'image': str(this_object.image).split('/')[-1],
        'image_path': '/media/' + str(this_object.image),
    })
